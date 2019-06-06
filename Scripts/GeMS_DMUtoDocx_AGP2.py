"""
Translates DMU table in NCGMP09-style geodatabase into a fully formatted
Microsoft Word .docx file.

Assumes formatting and style names in USGS Pubs template MapManuscript_v1-0_04-11.dotx

Arguments
    Input geodatabase
    Output workspace
    Output filename (if it doesn't end in .docx, .docx will be appended)
    UseMapUnitForUnitLabl (Boolean, either 'true' or 'false')
    
"""
# 4 June 2019: Edited to work with Python 3 in ArcGIS Pro - Evan Thoms
#   This version uses the new docx module for Python 3
#   https://python-docx.readthedocs.io/en/latest/
#   Install by opening the ArcGIS Pro Python Command Prompt so that you are starting in
#   the arcgispro-py3 conda environment.
#   Run 'conda install -c conda-forge python-docx'
#
#   Doesn't do all of the HTML conversion that Ralph's original does. So far it checks for:
#   <br> - line break/paragraph
#   <p> - paragraph, closing tag is not required and ignored if present
#   <b> - bold
#   <i> - italic
#   <sup> or <sub> super or subscript
#   <span  style="font-family: FGDCGeoAge"> write the enclosed text in FGDCGeoAge font, unit labels
#   also checks for non-printing line breaks if text with paragraphs is pasted from Word.
#
#   Consider updating with an HTML parse like BeautifulSoup. Although, it wasn't clear to me if bs4 is a 
#   solution for parsing text that is mostly un-formatted text with a few tags thrown in.

import sys, arcpy
import re
from pathlib import Path
from GeMS_utilityFunctions import *
import docx

versionString = 'GeMS_DMUtoDocx_AGP2.py, version of 5 June 2019'
addMsgAndPrint( versionString )

debug = False
debug2 = False

emDash = "\u2014"

startTags = []
endTags = []
tags = ['b','p','i','g','ul','sup','sub']
# bold, italic, group, unordered list, supercript, subscript
for tag in tags:
    startTags.append(r'<{}>'.format(tag))
    endTags.append(r'</{}>'.format(tag))
    
# span elements are special in that they contain attributes
# with this version, we'll only look for spans that are changing
# the font for unit labels
startTags.append('<span style="font-family: FGDCGeoAge">')
endTags.append('</span>')

def isNotBlank(thing):
    if thing != '' and thing != None:
        return True
    else:
        return False
    
def isKnownStyle(pStyle):
    for n in ['DMUHeadnote', 'DMU-Heading', 'DMUUnit']:
        if n in pStyle:
            return True
        else:
            return False

def notNullText(txt):
    if txt == '#null' or txt == None or txt == '#Null' or txt == '#' or txt == '' or len(txt.split()) == 0:
        return False
    else:
        return True
        
def add_formatting(paragraph, ptext):
    '''<paragraph> is a docx paragraph
   <ptext> is some text to be checked for HTML formatting tags'''
    # check to see if any tags appear in the text and, if so
    # collect a list of which ones
    res = [n for n in startTags if(n in ptext)]
    if res:
        # make a sorted list of the first to the last occurrence of all
        # startTags found in this text
        find_pos = []
        addMsgAndPrint('  Applying Word formatting for the following HTML tags:')
        addMsgAndPrint('  {}'.format(', '.join(res)))
        for st in res:
            et = endTags[startTags.index(st)]
            f_all = re.finditer(st, ptext)
            for match in f_all:
                begin = match.start()
                end = ptext.find(et, begin)
                find_pos.append([begin, end, st])
        find_pos.sort()

        # look for untagged text before the first tag
        before = ptext[0:find_pos[0][0]]
        if before:
            paragraph.add_run(before)
            
        # start slicing the rest
        for i, item in enumerate(find_pos):
            # the slice indices were stored during the creation of find_pos
            start = item[0] + len(item[2])
            end = item[1]
            tagged_text = ptext[start:end]
            
            # don't know a more elegant way to select the case for each of the tags
            # for which we are looking
            if item[2] == '<b>':
                paragraph.add_run(tagged_text).bold = True
            elif item[2] == '<i>':
                paragraph.add_run(tagged_text).italic = True
            elif item[2] == '<sup>':
                paragraph.add_run(tagged_text).font.superscript = True
            elif item[2] == '<sub>':
                paragraph.add_run(tagged_text).font.subscript = True
            elif item[2] == '<span style="font-family: FGDCGeoAge">':
                paragraph.add_run(tagged_text, 'DMUUnitLabeltypestyle')
    
            # find the untagged text between this tag and the next tag 
            # or between this tag and the end of the paragraph text
            # first, find the closing tag for the current tag
            et = endTags[startTags.index(item[2])]
            # and go len(et) steps beyond it to get to the start of the next slice
            start_n = end + len(et)
            # and find the index of the end of the next slice. As long as we
            # have not just evaluated the last tag in find_pos, the end index
            # is just before the next tag
            if i < len(find_pos) - 1:
                end_n = find_pos[i+1][0]
            else:
                # otherwise, the end index is the end of the paragraph text
                end_n = len(ptext)
               
            untagged_text = ptext[start_n:end_n]
            if untagged_text:
                paragraph.add_run(untagged_text)
    else:
        paragraph.add_run(ptext)

gdb = sys.argv[1]
outdir = Path(sys.argv[2])
outname = sys.argv[3]
if outname.lower()[-5:] != '.docx':
    outname = '{}.docx'.format(outname)
outDMUdocx = str(outdir.joinpath(outname))

if sys.argv[4] == 'true':
    useMapUnitForUnitLabl = True
else:
    useMapUnitForUnitLabl = False
    
if sys.argv[5] == 'true': # LMU only
    isLMU = True
else:
    isLMU = False

arcpy.env.workspace = gdb
script_parent = Path(sys.argv[0]).parent.parent
template_path = script_parent.joinpath("Resources", "MSWordDMUtemplate", "DMUtemplate.docx")
document = docx.Document(template_path)
document._body.clear_content()
lastParaWasHeading = False

"""
DMU has many rows
  Each row has content for 1 or more paragraphs. 1st paragraph has style 'row.ParagraphStyle'
  2nd and subsequent paragraphs have style 'DMUParagraph'
     Each paragraph is composed of one or more runs, each of which _may_ include
     markup tags

We sort DMU on HierarchyKey and then step through the rows, constructing rowtext w/ markup
  according to row.paragraphStyle.
We then divide the newly-built rowtext into paragraphs.
For each paragraph, we 

"""

addMsgAndPrint('Getting DMU rows and creating output paragraphs')
fields = ['mapunit', 'label', 'name', 'age', 'description', 'paragraphstyle', 'hierarchykey']
#          0          1        2       3      4              5                 6
sqlclause = (None, "ORDER by HierarchyKey ASC")
dmuRows = arcpy.da.SearchCursor('DescriptionOfMapUnits', fields, '#', '#', sqlclause)

# look for the case where we are making just a list of map units regardless of what 
# the title row of the table may be.
if isLMU:
    title = document.add_paragraph('LIST OF MAP UNITS', 'DMU-Heading1')
else:
    title = document.add_paragraph('DESCRIPTION OF MAP UNITS', 'DMU-Heading1')

# check the first row for a title and pass over it regardless because we named
# the document depending on the isLMU parameter   
first_row = dmuRows.next()
if first_row[2].lower in ['description of map units', 'list of map units']:
    row = dmuRows.next()

# check the first row for a headnote and add it only if we are building a full DMU
if first_row[5] == 'DMUHeadnote' and not isLMU:
    headnote = document.add_paragraph(first_row[4], 'DMUHeadnote')

for row in dmuRows:
    addMsgAndPrint('  {}: {}'.format(row[6], row[5]))
    if row[5].find('DMU-Heading') > -1:  # is a heading
        header_pr = document.add_paragraph(style=row[5])
        add_formatting(header_pr, row[2])
        if notNullText(row[4]):  # heading has headnote. Append heading as a paragraph
            headnote = document.add_paragraph(style='DMUHeadnote')
            add_formatting(headnote, row[4])
            #document.add_paragraph(row[4], 'DMUHeadnote')
                
    elif row[5].find('DMUUnit') > -1:  # is a unit
        # add an empty paragraph using ParagraphStyle
        unit = document.add_paragraph()
        if lastParaWasHeading:
            unit.style = 'DMUUnit11stafterheading'
        else:
            unit.style = row[5]
        
        # start this paragraph by adding the mapunit abbreviation, name, and age runs
        if not useMapUnitForUnitLabl and notNullText(row[1]):  
            abbrv = row[1]
        elif useMapUnitForUnitLabl and notNullText(row[0]):
            abbrv = row[0]
        abbrv = '{}\t'.format(abbrv)         
        if row[5][-1:] in ('4','5'):
            abbrv = '{}\t'.format(abbrv) # add second tab for DMUUnit4 and DMUUnit4
        unit.add_run(abbrv, 'DMUUnitLabeltypestyle')
        
        # check for name
        if isNotBlank(row[2]):
            unit.add_run(row[2], 'DMUUnitNameAgetypestyle')
        
        # check for age
        if isNotBlank(row[3]):
            unit.add_run('({})'.format(row[3], 'DMUUnitNameAgetypestyle'))
        
        # check for a description
        if isNotBlank(row[4]) and not isLMU:
            sep = ""
            desc_text = row[4]
            # try a couple ways to look for multiple paragraphs in the text
            # first, look for html break tags. If they are there, we'll trust 
            # that those are the only ones
            if "<br>" in desc_text:
                paras = desc_text.split("<br>")
            elif "<p>" in desc_text:
                paras = desc_text.split("<p>")
                paras = [t.replace("</p>", "") for t in paras]
            # otherwise, split on non-printing newline characters recognized by 
            # string.splitlines. Hopefully, this catches the rest
            else:
                # splitlines() is safe to use even if there are no lines in the text
                # we'll just get a list with one item
                paras = desc_text.splitlines()
            # in any case, add the first paragraph as a run or set of runs inside the current paragraph
            # first, add the emdash
            unit.add_run(emDash)
            
            addMsgAndPrint('  Evaluating paragraph 1')
            add_formatting(unit, paras[0])
            
            # and look for other paragraphs to add with DMUParagraph
            if len(paras) > 1 :
                for i in range(1, len(paras)):
                    addP = document.add_paragraph(style='DMUParagraph')
                    addMsgAndPrint('  Evaluating paragraph {}'.format(i+1))
                    add_formatting(addP, paras[i])
           
        if row[5].find('Head') > -1:
            lastParaWasHeading = True
        else:
            lastParaWasHeading = False
                    
    else: # Unrecognized paragraph style
        addMsgAndPrint('Do not recognize paragraph style {}'.format(row[5]))
        
addMsgAndPrint('    finished appending paragraphs')

if sys.argv[4] == 3:
    print(Null)
    pass
    
# Save our document
addMsgAndPrint('Saving to file {}'.format(outDMUdocx))
document.save(outDMUdocx)


